from extractor import extract_paragraphs
from translator import load_model_and_tokenizer, translate_batch
from docx import Document
from tqdm import tqdm

paragraphs = extract_paragraphs("docs/originals/CIJI.docx")
model, tokenizer = load_model_and_tokenizer()
translated_doc = Document()

batch_size = 8
batch = []

for para in tqdm(paragraphs, desc="Translating paragraphs"):
    text = para.strip()
    if not text:
        continue
    batch.append(text)

    if len(batch) == batch_size:
        translated_batch = translate_batch(batch, model, tokenizer)
        for t in translated_batch:
            translated_doc.add_paragraph(t)
        batch = []

if batch:
    translated_batch = translate_batch(batch, model, tokenizer)
    for t in translated_batch:
        translated_doc.add_paragraph(t)

translated_doc.save("docs/translate/Translate_CIJI.docx")
print("✅ Перевод завершён и сохранён в translated_book.docx")
